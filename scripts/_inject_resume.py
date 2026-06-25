"""重写项目二/三到 D:\\2026.docx,严格匹配项目一的字体/字号/颜色/双段 bullet 结构。

项目一规则:
- 项目标题:MicrosoftYaHei-Bold 20半磅(10pt) Bold #CC4125
- 时间:MicrosoftYaHei-Bold 20半磅(10pt) Bold #CC4125
- 子标题(角色):微软雅黑 18半磅(9pt) 不粗 #666666
- bullet 首段:MicrosoftYaHei-Bold 24半磅(12pt) Bold #CC4125
- bullet 续段:微软雅黑 18半磅(9pt) 不粗 #333333

每个 bullet = (首段, 续段) 元组。
"""
import copy
import shutil
import sys
from pathlib import Path

if sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOCX = Path(r"D:\2026.docx")
BACKUP = Path(r"D:\2026.injectv2.bak.docx")

# 字体规则
FONT_BOLD = "MicrosoftYaHei-Bold"
FONT_NORMAL = "微软雅黑"

# 颜色
COLOR_HIGHLIGHT = "CC4125"  # 红橙
COLOR_SUB = "666666"        # 中灰
COLOR_BODY = "333333"       # 深灰

# 字号(半磅)
SZ_TITLE = "20"     # 项目标题 / 时间 10pt
SZ_SUB = "18"       # 子标题 / 续段 9pt
SZ_HIGHLIGHT = "24" # bullet 首段 12pt


# ====== 项目二 ======
P2_HEADER_TITLE = "项目二  智慧社区 · AI 物业违规检测系统"
P2_SUB_TIME = "2025.12 – 2026.06"
P2_SUB_ROLE = "全栈 + AI 视觉算法 · 个人实战作品集"
P2_BULLETS = [
    (
        "·项目定位:社区场景 AI 视频安防系统原型,YOLO + 姿态分析 + LLM 自动复核三层架构",
        "覆盖违停 / 乱丢垃圾 / 宠物闯入 / 电瓶车进电梯 / 明火 / 翻越围栏 / 跌倒 7 类社区违规,V1→V2→V3 经三轮数据迭代优化",
    ),
    (
        "·数据闭环:从 24,632 → 19,396 → 36,459 张图,V3 模型 mAP@0.5 = 0.772",
        "融合 38 个 Roboflow 公开数据集 + VisDrone2019(无人机俯拍多目标街景),独立编写 5 个数据合并 / 转换 / 验证脚本",
    ),
    (
        "·模型部署:OpenVINO INT8 量化输出 11 MB(原始 22 MB,压缩 50%),CPU 单帧推理 25-40 ms",
        "AutoDL RTX 4090 训练 5.27 h;ONNX 通用格式导出 + RK3588 / BM1684X 边缘盒子转换脚本,部署链路打通",
    ),
    (
        "·后端架构:Flask 3 + SQLAlchemy + JWT + OpenVINO,12 个 API 蓝图 + 业务规则引擎",
        "覆盖检测 / 记录 / LLM / 摄像头 / 视频流 / 报表 / 审计 / 用户 / 居民端模块,SSE 实时事件推送",
    ),
    (
        "·前端架构:Vue 3 + Vite + Pinia + Element Plus + ECharts,15 个页面 + 9 个公共组件",
        "涵盖物业 PC 大屏(指挥中心 / 仪表盘 / 报表)、4 个检测入口(图像 / 视频 / 本机摄像头 / RTSP)、居民 H5 扫码端",
    ),
    (
        "·LLM 多服务商:抽象 OpenAI 兼容接口层,Qwen-VL / 智谱 GLM-4V / OpenAI / 豆包 4 家一键切换",
        "运行时主用切换实时生效;设计「YOLO 判定安全 → LLM 兜底校验」机制,降低姿态行为漏识风险",
    ),
    (
        "·代码规模:后端 Python ≈ 7,900 行 + 前端 Vue/JS ≈ 8,800 行,合计约 16,000 行",
        "完整 README + 边缘部署文档,单 CPU 即可推理,无需 GPU",
    ),
]

# ====== 项目三 ======
P3_HEADER_TITLE = "项目三  工地安防预警系统"
P3_SUB_TIME = "2025.04 – 2025.10"
P3_SUB_ROLE = "YOLO + LLM 双保险架构 · 全栈 + AI 算法 · 个人实战作品集 · 智慧社区方法论基础"
P3_BULLETS = [
    (
        "·项目定位:工地 AI 视频安防系统原型,YOLO 实时检测 + LLM 视觉二次研判双保险架构",
        "覆盖未戴安全帽 / 未穿反光衣 / 抽烟 / 打架 / 跌倒 / 高空作业未挂安全带 / 烟雾 / 火焰 8 类高风险违规",
    ),
    (
        "·核心架构 · YOLO + LLM 双保险:Qwen-VL / GLM-4V 视觉模型二次研判 confirmed=true/false",
        "降低 YOLO 误报;设计 prompt 引导 LLM 输出 JSON 结构化结果(画面描述 + 紧急度 + 整改建议)",
    ),
    (
        "·AI 工单文案:Qwen-Plus 自动生成「违规描述 + 处置建议」工单",
        "集成钉钉 / 企业微信 Webhook 推送通道,支持批量导出 PDF 报告(ReportLab)",
    ),
    (
        "·模型部署:YOLOv8s 训练 + OpenVINO INT8 量化 + ONNX 导出,CPU 单帧推理 25-30 ms",
        "无需 GPU,可部署到瑞芯微 / 算能边缘盒子",
    ),
    (
        "·RTSP 网络摄像头:OpenCV cv2.VideoCapture 直连海康 / 大华 RTSP 流",
        "MJPEG 推前端 + Canvas 叠加 AI 框,支持多路并发 + 断线自动重连",
    ),
    (
        "·整体架构:Flask 3 + Vue 3 + SQLAlchemy + JWT,11 个 API 蓝图 + 14 个前端页面",
        "后端 Python ≈ 4,300 行 + 前端 Vue/JS ≈ 6,200 行,合计约 10,500 行",
    ),
    (
        "·关键成果:沉淀 YOLO + LLM 双保险架构、业务规则引擎、RTSP 流接入等核心组件",
        "后续在「智慧社区」项目做领域迁移并扩展,从工地 B 端拓展到社区民生场景",
    ),
]


def set_run_style(rpr, font, size, bold, color):
    """设置 rPr 的字体/字号/加粗/颜色。"""
    # 字体(中英文都设)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:cs"), font)
    # 字号
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), size)
    # 加粗
    for tag in ("w:b", "w:bCs"):
        el = rpr.find(qn(tag))
        if bold:
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
        else:
            if el is not None:
                rpr.remove(el)
    # 颜色
    col = rpr.find(qn("w:color"))
    if col is None:
        col = OxmlElement("w:color")
        rpr.append(col)
    col.set(qn("w:val"), color)


def make_paragraph(after_p, text, font, size, bold, color):
    """在 after_p 之后插入新段落,带指定字体/字号/颜色;沿用 after_p 的段落属性 pPr。"""
    new_p = OxmlElement("w:p")
    ppr = after_p._element.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    r.append(rpr)
    set_run_style(rpr, font, size, bold, color)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_p.append(r)
    after_p._element.addnext(new_p)
    return Paragraph(new_p, after_p._parent)


def clear_para_set(paragraph, text, font, size, bold, color):
    """清空段落,设新文字 + 格式。"""
    for r in paragraph._element.findall(qn("w:r")):
        paragraph._element.remove(r)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    r.append(rpr)
    set_run_style(rpr, font, size, bold, color)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    paragraph._element.append(r)


def find_para(doc, prefix):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i, p
    return None, None


def delete_between(doc, start_idx, end_prefix):
    """删除 start_idx 之后到 end_prefix 之前的所有段落。"""
    end_idx = None
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith(end_prefix):
            end_idx = i
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    for i in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


def write_project(header_p, title, sub_time, sub_role, bullets):
    """改写头,后续插入时间 / 子标题 / 每个 bullet 的 (首段, 续段)。"""
    # 头:MicrosoftYaHei-Bold 20半磅 Bold #CC4125
    clear_para_set(header_p, title, FONT_BOLD, SZ_TITLE, True, COLOR_HIGHLIGHT)
    # 时间:同头
    last = make_paragraph(header_p, sub_time, FONT_BOLD, SZ_TITLE, True, COLOR_HIGHLIGHT)
    # 子标题:微软雅黑 18 不粗 #666666
    last = make_paragraph(last, sub_role, FONT_NORMAL, SZ_SUB, False, COLOR_SUB)
    # 每个 bullet 写 2 段
    for highlight, body in bullets:
        last = make_paragraph(last, highlight, FONT_BOLD, SZ_HIGHLIGHT, True, COLOR_HIGHLIGHT)
        if body:
            last = make_paragraph(last, body, FONT_NORMAL, SZ_SUB, False, COLOR_BODY)
    return last


def main():
    if not DOCX.exists():
        print(f"❌ 找不到 {DOCX}")
        sys.exit(1)
    shutil.copy2(DOCX, BACKUP)
    print(f"📦 备份: {BACKUP}")
    doc = Document(str(DOCX))

    # 项目二
    i2, p2 = find_para(doc, "项目二")
    if p2 is None:
        print("❌ 找不到「项目二」")
        sys.exit(1)
    print(f"📝 项目二 在第 {i2} 段")
    delete_between(doc, i2, "项目三")
    write_project(p2, P2_HEADER_TITLE, P2_SUB_TIME, P2_SUB_ROLE, P2_BULLETS)

    # 项目三
    i3, p3 = find_para(doc, "项目三")
    if p3 is None:
        print("❌ 找不到「项目三」")
        sys.exit(1)
    print(f"📝 项目三 在第 {i3} 段")
    delete_between(doc, i3, "自我评价")
    write_project(p3, P3_HEADER_TITLE, P3_SUB_TIME, P3_SUB_ROLE, P3_BULLETS)

    # 顺手替换"对标瑞幸咖啡小程序"
    OLD = "对标瑞幸咖啡小程序"
    NEW = "参考主流连锁咖啡品牌点单 + 会员模式"
    n = 0
    for p in doc.paragraphs:
        if OLD in p.text:
            # 用全文替换跨 run
            full = "".join(r.text for r in p.runs)
            new_full = full.replace(OLD, NEW)
            if p.runs:
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
                n += 1
    if n:
        print(f"📝 替换「{OLD}」→「{NEW}」共 {n} 处")

    doc.save(str(DOCX))
    print(f"\n✅ 写入完成 {DOCX}")
    print(f"   备份: {BACKUP}")


if __name__ == "__main__":
    main()
