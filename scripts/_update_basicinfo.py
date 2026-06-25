"""更新 2026.docx 基本信息块:
- 删除: 民族 / 婚姻 / 籍贯 / 政治面貌 / 国籍 (5 项)
- 修改: 期望薪资 >3K → 期望薪资 7K+
- 新增: 手机 / 邮箱 / GitHub (3 项,放在原删除位置上,不打乱布局)

策略: 用段落替换 + 倒序删除,保留原 run 的字体格式(rPr)。
"""
import copy
import shutil
import sys
from pathlib import Path

if sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path(r"D:\2026.docx")
BACKUP = Path(r"D:\2026.basicinfo.bak.docx")

# 用户提供
PHONE = "17760300240"
EMAIL = "3066974086@qq.com"
GITHUB = "github.com/[待补充用户名]"

# 字段替换映射:原前缀 → (新字段名, 新值) 或 None=删除
FIELD_MAPPING = {
    "民族":     ("手机", PHONE),
    "婚姻":     ("邮箱", EMAIL),
    "期望薪资":  ("期望薪资", "7K+"),
    "性别":     None,           # 保留原样
    "籍贯":     ("GitHub", GITHUB),
    "学历":     None,           # 保留
    "年龄":     None,           # 保留
    "政治面貌":  "DELETE",       # 删除
    "国籍":     "DELETE",       # 删除
}


def replace_runs(paragraph, new_field_name, new_field_value):
    """段落有 2 个 run(字段名 + 值),分别替换文字,保留 rPr 格式。"""
    runs = paragraph.runs
    if len(runs) >= 2:
        # 复制原格式
        runs[0].text = new_field_name
        runs[1].text = new_field_value
        # 把多余的 run 清掉
        for r in runs[2:]:
            r.text = ""
    else:
        # fallback: 整段重写
        for r in list(paragraph._element.findall(qn("w:r"))):
            paragraph._element.remove(r)
        # 用第一个原 run 的 rPr(如果有)
        new_r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_field_name + new_field_value
        t.set(qn("xml:space"), "preserve")
        new_r.append(t)
        paragraph._element.append(new_r)


def main():
    shutil.copy2(DOCX, BACKUP)
    print(f"📦 备份: {BACKUP}")
    doc = Document(str(DOCX))

    # 找"▍基本信息"和"▍教育经历"之间的段落
    info_start, info_end = None, None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("▍基本信息"):
            info_start = i
        elif info_start is not None and t.startswith("▍教育经历"):
            info_end = i
            break
    if info_start is None or info_end is None:
        print("❌ 找不到基本信息块")
        sys.exit(1)
    print(f"  基本信息块: 段 {info_start+1} ~ {info_end-1}")

    # 收集要删除的段落,倒序删除
    to_delete = []
    for i in range(info_start + 1, info_end):
        p = doc.paragraphs[i]
        text = p.text.strip()
        matched = False
        for prefix, action in FIELD_MAPPING.items():
            if text.startswith(prefix):
                matched = True
                if action == "DELETE":
                    to_delete.append(i)
                    print(f"  [{i}] {text:30s} → 删除")
                elif action is None:
                    print(f"  [{i}] {text:30s} → 保留")
                else:
                    new_name, new_val = action
                    print(f"  [{i}] {text:30s} → {new_name}{new_val}")
                    replace_runs(p, new_name, new_val)
                break
        if not matched:
            print(f"  [{i}] {text:30s} → 未匹配,跳过")

    # 倒序删
    for i in sorted(to_delete, reverse=True):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    doc.save(str(DOCX))
    print(f"\n✅ 已保存 {DOCX}")
    print(f"   备份: {BACKUP}")


if __name__ == "__main__":
    main()
