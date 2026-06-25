"""统一 D:\\2026.docx 字体风格(不美化,只规整):
- 正文段落:微软雅黑,18 半磅(9pt)
- ▍二级标题(基本信息/教育经历/项目经历/自我评价/技能与证书):微软雅黑 Bold,22 半磅(11pt)
- 项目标题(项目一/二/三/姓名后下方):微软雅黑 Bold,20 半磅(10pt)
- 姓名:保留 48 半磅(24pt)Bold
- 教育院校行(四川工业..):微软雅黑 Bold,20 半磅(10pt)
- 全部移除 Consolas / 新宋体 / 宋体 / Arial 等"杂"字体

也替换项目一里的"对标瑞幸咖啡小程序"为更得体的措辞。
"""
import copy
import shutil
import sys
from pathlib import Path

if sys.stdout.encoding.lower() != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path(r"D:\2026.docx")
BACKUP = Path(r"D:\2026.unifyfont.bak.docx")

NORMAL_SIZE = "18"   # 9pt
H2_SIZE = "22"       # 11pt  ▍标题
H3_SIZE = "20"       # 10pt  项目标题 / 学校行 / 求职意向
FONT_ZH = "微软雅黑"
FONT_EN = "微软雅黑"


def set_rfonts(rpr, font_zh=FONT_ZH, font_en=FONT_EN):
    """强制设置 rPr 的字体(ascii/eastAsia/hAnsi/cs 全部)。"""
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)
    rfonts.set(qn("w:eastAsia"), font_zh)
    rfonts.set(qn("w:cs"), font_en)


def set_size(rpr, half_pt):
    """w:sz 是半磅,例 18 = 9pt。"""
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), half_pt)
    # 同步 szCs (复杂字体)
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        rpr.append(szcs)
    szcs.set(qn("w:val"), half_pt)


def set_bold(rpr, bold=True):
    b = rpr.find(qn("w:b"))
    if bold:
        if b is None:
            b = OxmlElement("w:b")
            rpr.append(b)
    else:
        if b is not None:
            rpr.remove(b)
    # 同步 bCs
    bcs = rpr.find(qn("w:bCs"))
    if bold:
        if bcs is None:
            bcs = OxmlElement("w:bCs")
            rpr.append(bcs)
    else:
        if bcs is not None:
            rpr.remove(bcs)


def style_paragraph(paragraph, size_half_pt=NORMAL_SIZE, bold=False):
    """统一段落里所有 run 的字体、字号、加粗。"""
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._r.insert(0, rpr)
        set_rfonts(rpr)
        set_size(rpr, size_half_pt)
        set_bold(rpr, bold)


def classify(text):
    """根据段落文字判断段落类型,返回 (size_half_pt, bold)。"""
    t = text.strip()
    if not t:
        return None
    # ▍二级标题:基本信息 / 教育经历 / 项目经历 / 自我评价 / 技能与证书
    if t.startswith("▍"):
        return (H2_SIZE, True)
    # 项目一/二/三标题
    if t.startswith("项目一") or t.startswith("项目二") or t.startswith("项目三"):
        return (H3_SIZE, True)
    # 求职意向(整段)
    if t.startswith("求职意向"):
        return (H3_SIZE, True)
    # 学校行
    if "四川工业科技学院" in t and "本科" in t:
        return (H3_SIZE, True)
    # 自我评价(纯文字标题段)
    if t == "自我评价":
        return (H2_SIZE, True)
    # 其他:正文
    return (NORMAL_SIZE, False)


def replace_text_in_runs(paragraph, old, new):
    """跨 run 的"对标瑞幸"替换:把所有 run 文字合到一起替换,再分配回去。"""
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # 简单做法:把第一个 run 设为新文本,其他清空
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def main():
    shutil.copy2(DOCX, BACKUP)
    print(f"📦 备份: {BACKUP}")
    doc = Document(str(DOCX))

    n_changed = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        # 姓名段不动(48 半磅 Bold 保持)
        if t == "胡礼超":
            continue
        cls = classify(t)
        if cls is None:
            continue
        size, bold = cls
        style_paragraph(p, size_half_pt=size, bold=bold)
        n_changed += 1

    # 替换"对标瑞幸咖啡小程序"为更得体的措辞
    OLD = "对标瑞幸咖啡小程序"
    NEW = "参考主流连锁咖啡品牌点单 + 会员模式"
    replaced = 0
    for p in doc.paragraphs:
        if replace_text_in_runs(p, OLD, NEW):
            replaced += 1
            # 替换后段落字体可能因合并 run 丢失,重设
            cls = classify(p.text)
            if cls:
                style_paragraph(p, *cls)

    doc.save(str(DOCX))
    print(f"✅ 字体统一: {n_changed} 段")
    print(f"✅ 文案替换: 「{OLD}」 → 「{NEW}」 共 {replaced} 处")
    print(f"   备份: {BACKUP}")


if __name__ == "__main__":
    main()
